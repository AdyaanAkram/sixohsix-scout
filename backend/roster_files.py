"""Multi-format roster file parsing — CSV, Excel (.xlsx), Word tables (.docx).

Both importers (org Players import and event roster import) call rows_from_upload
and get back (fieldnames, rows) exactly as csv.DictReader would produce, so the
existing header-mapping pipeline stays the single source of truth for what any
column means. Google Docs/Sheets users export via File → Download → .docx/.xlsx.
"""
from __future__ import annotations

import csv
import io

from fastapi import HTTPException, UploadFile

ROSTER_FILE_EXTS = (".csv", ".xlsx", ".docx")
MAX_ROSTER_FILE_BYTES = 5 * 1024 * 1024


def _cell_text(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))  # Excel loves turning 2032 into 2032.0
    return str(value).strip()


def _rows_from_csv(raw: bytes):
    try:
        text = raw.decode("utf-8-sig")
    except UnicodeDecodeError:
        text = raw.decode("latin-1")
    reader = csv.DictReader(io.StringIO(text))
    return reader.fieldnames, list(reader)


def _rows_from_xlsx(raw: bytes):
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    except Exception:
        raise HTTPException(status_code=400, detail="This Excel file could not be read. Re-save it as .xlsx and try again.")
    ws = wb.worksheets[0]
    it = ws.iter_rows(values_only=True)
    header = next(it, None)
    if not header:
        raise HTTPException(status_code=400, detail="The first sheet is empty.")
    fieldnames = [_cell_text(h) for h in header]
    rows = []
    for r in it:
        row = {fieldnames[i]: _cell_text(r[i]) if i < len(r) else "" for i in range(len(fieldnames))}
        if any(v for v in row.values()):
            rows.append(row)
    wb.close()
    return [f for f in fieldnames if f], rows


def _rows_from_docx(raw: bytes):
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw))
    except Exception:
        raise HTTPException(status_code=400, detail="This Word file could not be read. Re-save it as .docx and try again.")
    if not doc.tables:
        raise HTTPException(
            status_code=400,
            detail="No table found in the Word document. Put the roster in a table with a header row "
                   "(in Google Docs: Insert → Table), or export as CSV/Excel instead.")
    table = doc.tables[0]
    header = [c.text.strip() for c in table.rows[0].cells]
    fieldnames = [h for h in header if h]
    rows = []
    for tr in table.rows[1:]:
        cells = [c.text.strip() for c in tr.cells]
        row = {header[i]: cells[i] if i < len(cells) else "" for i in range(len(header)) if header[i]}
        if any(v for v in row.values()):
            rows.append(row)
    return fieldnames, rows


async def rows_from_upload(file: UploadFile):
    """(fieldnames, rows) from a roster upload in any supported format."""
    name = (file.filename or "").lower()
    if not name.endswith(ROSTER_FILE_EXTS):
        raise HTTPException(
            status_code=400,
            detail="Please upload a .csv, .xlsx (Excel/Google Sheets) or .docx (Word/Google Docs) file.")
    raw = await file.read()
    if len(raw) > MAX_ROSTER_FILE_BYTES:
        raise HTTPException(status_code=400, detail="File is too large (max 5 MB).")
    if name.endswith(".csv"):
        fieldnames, rows = _rows_from_csv(raw)
    elif name.endswith(".xlsx"):
        fieldnames, rows = _rows_from_xlsx(raw)
    else:
        fieldnames, rows = _rows_from_docx(raw)
    if not fieldnames:
        raise HTTPException(status_code=400, detail="The file appears to have no header row.")
    return fieldnames, rows
